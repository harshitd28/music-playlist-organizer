#!/usr/bin/env python3
"""
Convert PROJECT_REPORT.md to Word document (.docx)
"""

import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def process_bold_text(para, text):
    """Process text with bold markers"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    current_table = None
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        original_line = line
        line = line.rstrip()
        
        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    para = doc.add_paragraph('\n'.join(code_lines))
                    para.style = 'Intense Quote'
                    run = para.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines (but allow some spacing)
        if not line.strip():
            i += 1
            continue
        
        # Heading level 1
        if line.startswith('# ') and not line.startswith('##'):
            text = line[2:].strip()
            para = doc.add_heading(text, level=1)
            i += 1
        
        # Heading level 2
        elif line.startswith('## ') and not line.startswith('###'):
            text = line[3:].strip()
            if text == '---':
                i += 1
                continue
            doc.add_heading(text, level=2)
            i += 1
        
        # Heading level 3
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
        
        # Heading level 4
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
            i += 1
        
        # Horizontal rule
        elif line.strip() == '---':
            i += 1
            continue
        
        # Table detection
        elif '|' in line and line.strip().startswith('|'):
            # Check if it's a separator line
            if re.match(r'^\|[\s\-:]+\|', line):
                current_table = None
                i += 1
                continue
            
            # Parse table row
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells from edges
            
            if len(cells) > 1:
                if current_table is None:
                    current_table = doc.add_table(rows=1, cols=len(cells))
                    current_table.style = 'Light Grid Accent 1'
                    # Make header row bold
                    for j, cell_text in enumerate(cells):
                        cell = current_table.rows[0].cells[j]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                else:
                    row = current_table.add_row()
                    for j, cell_text in enumerate(cells):
                        if j < len(row.cells):
                            row.cells[j].text = cell_text
            i += 1
        
        # Bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            para = doc.add_paragraph(text, style='List Bullet')
            if '**' in text:
                para.clear()
                process_bold_text(para, text)
            i += 1
        
        # Numbered list
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            para = doc.add_paragraph(text, style='List Number')
            if '**' in text:
                para.clear()
                process_bold_text(para, text)
            i += 1
        
        # Regular paragraph with potential formatting
        else:
            if current_table is not None:
                current_table = None
            
            text = line.strip()
            
            # Create paragraph
            para = doc.add_paragraph()
            
            # Process bold text
            if '**' in text:
                process_bold_text(para, text)
            else:
                # Clean up other markdown
                text = re.sub(r'`([^`]+)`', r'\1', text)  # Inline code
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
                para.add_run(text)
            
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")
    print(f"📄 Word document saved as: {docx_file}")

if __name__ == '__main__':
    md_file = 'PROJECT_REPORT.md'
    docx_file = 'PROJECT_REPORT.docx'
    parse_markdown_to_docx(md_file, docx_file)
